"""
Seed onboarding demo data:
- 5 Contract DOCX templates (one per work_type) with placeholders
- 1 Job Offer DOCX template
- Sample Job Description templates for common roles
- 1 Employee Invite for demo onboarding
"""
import uuid
from io import BytesIO
from datetime import timedelta

from django.core.management.base import BaseCommand
from django.core.files.base import ContentFile
from django.utils import timezone
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from users.models import Organization, UserProfile
from hr.models import DocumentTemplate, EmployeeInvite


def _make_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return h


def _add_para(doc, text, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p


def _add_field_row(doc, label, value):
    p = doc.add_paragraph()
    run_label = p.add_run(f'{label}: ')
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_val = p.add_run(value)
    run_val.font.size = Pt(10)
    return p


def create_contract_docx(work_type_label):
    """Create a contract DOCX template for a specific work type."""
    doc = Document()

    # Title
    title = doc.add_heading('Employment Agreement', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    _add_para(doc, f'{work_type_label} Employment Contract', bold=True, size=12)
    doc.add_paragraph('')

    # Parties
    _make_heading(doc, '1. Parties', level=2)
    _add_field_row(doc, 'Employer', '{{company_name}}')
    _add_field_row(doc, 'Address', '{{company_address}}')
    _add_field_row(doc, 'Phone', '{{company_phone}}')
    _add_field_row(doc, 'Email', '{{company_email}}')
    _add_field_row(doc, 'IRD Number', '{{company_ird}}')
    doc.add_paragraph('')
    _add_field_row(doc, 'Employee', '{{employee_name}}')
    _add_field_row(doc, 'Email', '{{employee_email}}')

    # Position
    _make_heading(doc, '2. Position Details', level=2)
    _add_field_row(doc, 'Job Title', '{{job_title}}')
    _add_field_row(doc, 'Employment Type', '{{work_type}}')
    _add_field_row(doc, 'Start Date', '{{start_date}}')

    # Remuneration
    _make_heading(doc, '3. Remuneration', level=2)
    _add_field_row(doc, 'Hourly Rate', '${{hourly_rate}} per hour (before tax)')
    _add_para(doc, 'Pay will be made fortnightly by direct credit to the employee\'s nominated bank account.')

    # Hours
    _make_heading(doc, '4. Hours of Work', level=2)
    if work_type_label == 'Full Time':
        _add_para(doc, 'The employee will work a minimum of 30 hours per week. The employer and employee will agree on the days and hours of work.')
    elif work_type_label == 'Part Time':
        _add_para(doc, 'The employee will work agreed hours each week (less than 30 hours). Specific days and hours will be confirmed by the roster.')
    elif work_type_label == 'Casual':
        _add_para(doc, 'Work will be offered on an as-needed basis. The employee is not obligated to accept any particular offer of work. Each engagement is a separate period of employment.')
    elif work_type_label == 'Salary':
        _add_para(doc, 'The employee is expected to work the hours reasonably necessary to fulfil the requirements of the role. Standard hours are 40 hours per week.')
    elif work_type_label == 'Visa Full Time':
        _add_para(doc, 'The employee will work a minimum of 30 hours per week, subject to the conditions of their work visa.')

    # Leave
    _make_heading(doc, '5. Leave Entitlements', level=2)
    _add_para(doc, 'Annual Leave: After 12 months of continuous employment, the employee is entitled to 4 weeks paid annual leave per year.')
    _add_para(doc, 'Sick Leave: After 6 months of continuous employment, the employee is entitled to 10 days paid sick leave per year.')
    _add_para(doc, 'Public Holidays: The employee is entitled to public holidays in accordance with the Holidays Act 2003.')

    # Trial Period
    _make_heading(doc, '6. Trial Period', level=2)
    _add_para(doc, 'This employment is subject to a 90-day trial period commencing from the start date. During this period, the employer may dismiss the employee and the employee may not bring a personal grievance on the grounds of unjustified dismissal.')

    # Termination
    _make_heading(doc, '7. Termination', level=2)
    _add_para(doc, 'Either party may terminate this agreement by giving 2 weeks written notice. The employer may terminate without notice in cases of serious misconduct.')

    # Signatures
    _make_heading(doc, '8. Signatures', level=2)
    doc.add_paragraph('')
    _add_para(doc, 'Employer Signature: ___________________________    Date: ___________')
    doc.add_paragraph('')
    _add_para(doc, 'Employee Signature: ___________________________    Date: ___________')
    doc.add_paragraph('')
    _add_para(doc, 'Employee Name: {{employee_name}}')

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def create_job_offer_docx():
    """Create a Job Offer DOCX template."""
    doc = Document()

    title = doc.add_heading('Job Offer Letter', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    _add_field_row(doc, 'Date', '{{start_date}}')
    doc.add_paragraph('')

    _add_para(doc, 'Dear {{employee_name}},', size=11)
    doc.add_paragraph('')

    _add_para(doc,
        'We are pleased to offer you the position of {{job_title}} at {{company_name}}. '
        'We believe your skills and experience will be a valuable addition to our team.',
        size=11)
    doc.add_paragraph('')

    _make_heading(doc, 'Offer Details', level=2)
    _add_field_row(doc, 'Position', '{{job_title}}')
    _add_field_row(doc, 'Employment Type', '{{work_type}}')
    _add_field_row(doc, 'Hourly Rate', '${{hourly_rate}} per hour')
    _add_field_row(doc, 'Start Date', '{{start_date}}')
    _add_field_row(doc, 'Location', '{{company_address}}')
    doc.add_paragraph('')

    _add_para(doc,
        'This offer is conditional upon satisfactory completion of a 90-day trial period '
        'and verification of your right to work in New Zealand.',
        size=11)
    doc.add_paragraph('')

    _add_para(doc,
        'Please confirm your acceptance by signing and returning this letter.',
        size=11)
    doc.add_paragraph('')

    _add_para(doc, 'Sincerely,')
    _add_para(doc, '{{company_name}}')
    _add_para(doc, '{{company_phone}} | {{company_email}}')
    doc.add_paragraph('')
    doc.add_paragraph('')

    _add_para(doc, 'I, {{employee_name}}, accept this offer of employment.', bold=True)
    doc.add_paragraph('')
    _add_para(doc, 'Signature: ___________________________    Date: ___________')

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def create_job_description_docx(job_title_label, duties):
    """Create a Job Description DOCX template."""
    doc = Document()

    title = doc.add_heading('Job Description', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    _make_heading(doc, job_title_label, level=1)
    doc.add_paragraph('')

    _add_field_row(doc, 'Company', '{{company_name}}')
    _add_field_row(doc, 'Location', '{{company_address}}')
    _add_field_row(doc, 'Reports To', 'Store Manager')
    _add_field_row(doc, 'Employment Type', '{{work_type}}')
    doc.add_paragraph('')

    _make_heading(doc, 'Key Responsibilities', level=2)
    for duty in duties:
        p = doc.add_paragraph(duty, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_paragraph('')
    _make_heading(doc, 'Requirements', level=2)
    _add_para(doc, 'Must have a positive attitude and strong work ethic.')
    _add_para(doc, 'Ability to work in a fast-paced team environment.')
    _add_para(doc, 'Flexible availability including weekends and public holidays.')

    doc.add_paragraph('')
    _add_para(doc, 'Employee: {{employee_name}}', bold=True)
    _add_para(doc, 'Acknowledged on: {{start_date}}')

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# Job descriptions per role
JOB_DUTIES = {
    'BARISTA': [
        'Prepare and serve espresso-based beverages to company standards',
        'Operate and maintain coffee machine and grinders',
        'Take customer orders and process payments accurately',
        'Maintain cleanliness and hygiene of the coffee station',
        'Manage stock levels of coffee beans, milk, and supplies',
        'Provide friendly and efficient customer service',
    ],
    'ALL_ROUNDER': [
        'Assist with food preparation and service as required',
        'Operate the point-of-sale system and process transactions',
        'Maintain cleanliness of all customer and staff areas',
        'Restock supplies and manage inventory',
        'Support all team members across different stations',
        'Follow food safety and hygiene procedures at all times',
    ],
    'SERVER': [
        'Greet and seat customers in a friendly manner',
        'Present menus and provide recommendations',
        'Take customer orders accurately and relay to kitchen',
        'Serve food and beverages to tables',
        'Process payments and handle cash/card transactions',
        'Ensure tables are clean and properly set',
    ],
    'CASHIER': [
        'Operate the point-of-sale system accurately',
        'Process cash, card, and other payment methods',
        'Greet customers and provide excellent service',
        'Reconcile cash register at end of shift',
        'Maintain cleanliness of the counter area',
        'Assist with front-of-house duties as needed',
    ],
    'KITCHEN_HAND': [
        'Wash and sanitise dishes, utensils, and kitchen equipment',
        'Assist chefs with basic food preparation tasks',
        'Maintain cleanliness and organisation of the kitchen',
        'Receive and store deliveries correctly',
        'Dispose of waste and recycling appropriately',
        'Follow all food safety and hygiene procedures',
    ],
}


class Command(BaseCommand):
    help = 'Seed onboarding demo: DOCX templates + employee invite'

    def handle(self, *args, **options):
        org = Organization.objects.first()
        if not org:
            self.stderr.write('No organization found.')
            return

        manager = UserProfile.objects.filter(
            organization=org,
            role__in=['MANAGER', 'SENIOR_MANAGER', 'CEO'],
        ).first()
        if not manager:
            self.stderr.write('No manager found.')
            return

        self.stdout.write(f'Using org: {org.name} (id={org.id})')
        self.stdout.write(f'Using manager: {manager.user.username}')

        # Clear old templates for this org
        old_count = DocumentTemplate.objects.filter(organization=org).count()
        DocumentTemplate.objects.filter(organization=org).delete()
        self.stdout.write(f'Cleared {old_count} old templates')

        # ── Contract Templates (5 work types) ──
        work_types = [
            ('FULL_TIME', 'Full Time'),
            ('PART_TIME', 'Part Time'),
            ('CASUAL', 'Casual'),
            ('SALARY', 'Salary'),
            ('VISA_FULL_TIME', 'Visa Full Time'),
        ]
        for wt_key, wt_label in work_types:
            buf = create_contract_docx(wt_label)
            tmpl = DocumentTemplate.objects.create(
                organization=org,
                document_type='CONTRACT',
                work_type=wt_key,
                title=f'{wt_label} Employment Agreement',
            )
            tmpl.file.save(
                f'contract_{wt_key.lower()}.docx',
                ContentFile(buf.read()),
                save=True,
            )
            self.stdout.write(f'  + Contract: {wt_label}')

        # ── Job Offer Template ──
        buf = create_job_offer_docx()
        tmpl = DocumentTemplate.objects.create(
            organization=org,
            document_type='JOB_OFFER',
            title='Job Offer Letter',
        )
        tmpl.file.save('job_offer.docx', ContentFile(buf.read()), save=True)
        self.stdout.write('  + Job Offer')

        # ── Job Description Templates ──
        for jt_key, duties in JOB_DUTIES.items():
            label = jt_key.replace('_', ' ').title()
            buf = create_job_description_docx(label, duties)
            tmpl = DocumentTemplate.objects.create(
                organization=org,
                document_type='JOB_DESCRIPTION',
                job_title=jt_key,
                title=f'{label} Job Description',
            )
            tmpl.file.save(
                f'jd_{jt_key.lower()}.docx',
                ContentFile(buf.read()),
                save=True,
            )
            self.stdout.write(f'  + JD: {label}')

        # ── Demo Employee Invite ──
        # Delete old pending invites
        EmployeeInvite.objects.filter(
            organization=org, status='PENDING'
        ).delete()

        invite = EmployeeInvite.objects.create(
            organization=org,
            email='demo.employee@test.com',
            first_name='James',
            last_name='Kim',
            role='EMPLOYEE',
            job_title='BARISTA',
            work_type='FULL_TIME',
            hourly_rate=23.15,
            status='PENDING',
            invited_by=manager,
            expires_at=timezone.now() + timedelta(days=7),
        )

        self.stdout.write('')
        self.stdout.write(self.style.SUCCESS('=== Demo Onboarding Ready ==='))
        self.stdout.write(f'Invite Code: {invite.invite_code}')
        self.stdout.write(f'Employee: James Kim (demo.employee@test.com)')
        self.stdout.write(f'Position: Barista / Full Time / $23.15/hr')
        self.stdout.write(f'Accept URL: /accept-invite?code={invite.invite_code}')
        self.stdout.write('')
        self.stdout.write('To test: Accept the invite using the code above.')
