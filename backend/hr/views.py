import uuid
import base64
from rest_framework import viewsets, status, mixins, generics
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.views import APIView
from django.contrib.auth.models import User
from django.core.files.base import ContentFile
from django.conf import settings
from django.core.mail import send_mail
from django.db import transaction
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.utils import timezone
from datetime import timedelta
import logging

logger = logging.getLogger(__name__)

from utils.docx_filler import fill_docx_template
from utils.docx_to_pdf import convert_docx_to_pdf
from utils.pdf_signer import sign_pdf
from utils.pdf_zones import extract_and_clean_sign_zones, get_default_zones

from .models import (
    Onboarding, OnboardingTask, EmployeeDocument, ShiftTemplate, Roster, Timesheet, Task,
    EmployeeInvite, DocumentTemplate, IR330Declaration, TrainingModule, Inquiry, ResignationRequest,
    DisciplinaryRecord, PerformanceReview, WorkplaceAccident, EmployeeNote,
)
from .serializers import (
    OnboardingDetailSerializer, OnboardingListSerializer,
    OnboardingTaskSerializer, EmployeeDocumentSerializer,
    ShiftTemplateSerializer, RosterSerializer, TimesheetSerializer, TaskSerializer,
    EmployeeInviteSerializer, DocumentTemplateSerializer, IR330DeclarationSerializer,
    TrainingModuleSerializer, TeamMemberListSerializer, TeamMemberDetailSerializer,
    InquirySerializer, ResignationRequestSerializer,
    DisciplinaryRecordSerializer, PerformanceReviewSerializer,
    WorkplaceAccidentSerializer, EmployeeNoteSerializer,
)
from users.models import UserProfile, ROLE_CHOICES, JOB_TITLE_CHOICES, WORK_TYPE_CHOICES
from users.serializers import UserProfileDetailSerializer
from users.permissions import IsManager, IsEmployee
from users.filters import OrganizationFilterBackend, get_target_org
from payroll.models import Salary


class OnboardingViewSet(viewsets.ModelViewSet):
    """
    온보딩 프로세스 관리 ViewSet
    - list: 온보딩 목록
    - create: 신규 온보딩 생성
    - retrieve: 온보딩 상세 조회 (직원 본인도 접근 가능)
    - update/partial_update: 온보딩 수정
    - destroy: 온보딩 삭제
    - complete: 온보딩 완료 (직원 본인도 접근 가능)
    """
    queryset = Onboarding.objects.all()
    permission_classes = [IsAuthenticated, IsManager]
    filter_backends = [OrganizationFilterBackend]

    def get_permissions(self):
        if self.action in ['retrieve', 'complete']:
            return [IsAuthenticated()]
        return [IsAuthenticated(), IsManager()]

    def get_serializer_class(self):
        if self.action == 'retrieve':
            return OnboardingDetailSerializer
        return OnboardingListSerializer

    def get_queryset(self):
        """사용자의 조직에 해당하는 온보딩만 조회"""
        queryset = super().get_queryset()
        return queryset.select_related('employee__user', 'assigned_to__user', 'organization')

    def perform_create(self, serializer):
        """온보딩 생성 시 자동으로 현재 사용자 할당"""
        serializer.save()

    @action(detail=True, methods=['post'])
    def complete(self, request, pk=None):
        """온보딩 완료 표시"""
        onboarding = self.get_object()

        if not onboarding.all_tasks_completed:
            return Response(
                {'error': '모든 작업을 완료한 후 온보딩을 완료할 수 있습니다.'},
                status=status.HTTP_400_BAD_REQUEST
            )

        onboarding.status = 'COMPLETED'
        onboarding.completed_percentage = 100
        onboarding.completed_at = timezone.now()
        onboarding.save()

        return Response(
            {'message': '온보딩이 완료되었습니다.'},
            status=status.HTTP_200_OK
        )

    @action(detail=True, methods=['get'])
    def progress(self, request, pk=None):
        """온보딩 진행률 조회"""
        onboarding = self.get_object()
        total_tasks = onboarding.tasks.count()
        completed_tasks = onboarding.tasks.filter(is_completed=True).count()

        percentage = (completed_tasks / total_tasks * 100) if total_tasks > 0 else 0

        return Response({
            'total_tasks': total_tasks,
            'completed_tasks': completed_tasks,
            'percentage': int(percentage),
            'status': onboarding.status
        }, status=status.HTTP_200_OK)


class OnboardingTaskViewSet(viewsets.ModelViewSet):
    """
    온보딩 작업 항목 관리 ViewSet
    - list: 작업 목록
    - create: 신규 작업 생성
    - update/partial_update: 작업 수정 (직원 본인도 접근 가능 — 파일 업로드 등)
    - destroy: 작업 삭제
    - complete: 작업 완료 표시 (직원 본인도 접근 가능)
    """
    queryset = OnboardingTask.objects.all()
    serializer_class = OnboardingTaskSerializer
    permission_classes = [IsAuthenticated, IsManager]

    def get_permissions(self):
        if self.action in ['complete', 'incomplete', 'partial_update']:
            return [IsAuthenticated()]
        return [IsAuthenticated(), IsManager()]

    def get_queryset(self):
        """쿼리 파라미터로 온보딩 필터링"""
        queryset = super().get_queryset()
        onboarding_id = self.request.query_params.get('onboarding_id')

        if onboarding_id:
            queryset = queryset.filter(onboarding_id=onboarding_id)

        return queryset.select_related('assigned_to__user', 'onboarding__employee__user')

    @action(detail=True, methods=['post'])
    def complete(self, request, pk=None):
        """작업 완료 표시"""
        task = self.get_object()
        task.is_completed = True
        task.completed_at = timezone.now()
        task.save()

        # 온보딩 진행률 업데이트
        onboarding = task.onboarding
        total_tasks = onboarding.tasks.count()
        completed_tasks = onboarding.tasks.filter(is_completed=True).count()
        onboarding.completed_percentage = int(completed_tasks / total_tasks * 100) if total_tasks > 0 else 0
        onboarding.save()

        return Response(
            {'message': '작업이 완료되었습니다.'},
            status=status.HTTP_200_OK
        )

    @action(detail=True, methods=['post'])
    def incomplete(self, request, pk=None):
        """작업 완료 취소"""
        task = self.get_object()
        task.is_completed = False
        task.completed_at = None
        task.save()

        # 온보딩 진행률 업데이트
        onboarding = task.onboarding
        total_tasks = onboarding.tasks.count()
        completed_tasks = onboarding.tasks.filter(is_completed=True).count()
        onboarding.completed_percentage = int(completed_tasks / total_tasks * 100) if total_tasks > 0 else 0
        onboarding.save()

        return Response(
            {'message': '작업 완료가 취소되었습니다.'},
            status=status.HTTP_200_OK
        )


class EmployeeDocumentViewSet(viewsets.ModelViewSet):
    """
    직원 문서 관리 ViewSet
    - list: 문서 목록
    - create: 신규 문서 업로드
    - retrieve: 문서 상세 조회
    - destroy: 문서 삭제
    - sign: 문서 서명
    - download: 문서 다운로드
    """
    queryset = EmployeeDocument.objects.all()
    serializer_class = EmployeeDocumentSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        """쿼리 파라미터로 온보딩 필터링"""
        queryset = super().get_queryset()
        onboarding_id = self.request.query_params.get('onboarding_id')

        if onboarding_id:
            queryset = queryset.filter(onboarding_id=onboarding_id)

        return queryset.select_related('signed_by__user', 'uploaded_by__user', 'onboarding__employee__user')

    def perform_create(self, serializer):
        """문서 업로드 시 업로드 사용자 자동 할당"""
        serializer.save(uploaded_by=self.request.user.profile)

    @action(detail=True, methods=['post'])
    def sign(self, request, pk=None):
        """문서 서명 (서명 + 이니셜 이미지를 PDF에 합성)"""
        document = self.get_object()

        if document.is_signed:
            return Response(
                {'error': '이미 서명된 문서입니다.'},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Decode base64 images
        signature_bytes = None
        initials_bytes = None

        signature_data = request.data.get('signature')
        if signature_data and ';base64,' in signature_data:
            try:
                fmt, imgstr = signature_data.split(';base64,')
                signature_bytes = base64.b64decode(imgstr)
                ext = fmt.split('/')[-1] if '/' in fmt else 'png'
                document.signature.save(
                    f'sig_{document.id}.{ext}',
                    ContentFile(signature_bytes),
                    save=False,
                )
            except Exception:
                pass

        initials_data = request.data.get('initials')
        if initials_data and ';base64,' in initials_data:
            try:
                _, imgstr = initials_data.split(';base64,')
                initials_bytes = base64.b64decode(imgstr)
            except Exception:
                pass

        # Ensure PDF exists (convert DOCX → PDF if needed)
        if not document.pdf_file and document.file:
            if document.file.name.lower().endswith('.docx'):
                pdf_bytes = convert_docx_to_pdf(document.file)
                if pdf_bytes:
                    fname = f'{document.document_type.lower()}_{document.id}.pdf'
                    document.pdf_file.save(fname, ContentFile(pdf_bytes), save=False)

        # Overlay signature + initials onto the PDF
        if document.pdf_file and (signature_bytes or initials_bytes):
            try:
                document.pdf_file.open('rb')
                original_pdf = document.pdf_file.read()
                document.pdf_file.close()
                signed_pdf = sign_pdf(original_pdf, signature_bytes, initials_bytes, document.sign_zones)
                document.pdf_file.save(
                    document.pdf_file.name.split('/')[-1],
                    ContentFile(signed_pdf),
                    save=False,
                )
            except Exception:
                pass

        document.is_signed = True
        document.signed_at = timezone.now()
        document.signed_by = request.user.profile
        document.save()

        return Response(
            {'message': '문서가 서명되었습니다.'},
            status=status.HTTP_200_OK
        )

    @action(detail=True, methods=['get'])
    def preview(self, request, pk=None):
        """문서 PDF 미리보기 — DOCX는 자동 변환, PDF는 그대로 서빙"""
        document = self.get_object()

        # 1. 이미 캐시된 PDF가 있으면 바로 서빙
        if document.pdf_file:
            response = HttpResponse(
                document.pdf_file.read(), content_type='application/pdf',
            )
            response['Content-Disposition'] = f'inline; filename="{document.title}.pdf"'
            return response

        # 2. 원본이 PDF이면 바로 서빙
        if document.file and document.file.name.lower().endswith('.pdf'):
            response = HttpResponse(
                document.file.read(), content_type='application/pdf',
            )
            response['Content-Disposition'] = f'inline; filename="{document.title}.pdf"'
            return response

        # 3. DOCX → PDF 변환 후 캐시
        if document.file and document.file.name.lower().endswith('.docx'):
            pdf_bytes = convert_docx_to_pdf(document.file)
            if pdf_bytes:
                filename = f'{document.document_type.lower()}_{document.id}.pdf'
                document.pdf_file.save(filename, ContentFile(pdf_bytes), save=True)
                response = HttpResponse(pdf_bytes, content_type='application/pdf')
                response['Content-Disposition'] = f'inline; filename="{document.title}.pdf"'
                return response

        # 4. Fallback: 원본 파일 서빙
        if document.file:
            response = HttpResponse(document.file.read())
            response['Content-Disposition'] = f'inline; filename="{document.title}"'
            return response

        return Response({'error': 'No file found'}, status=404)

    @action(detail=True, methods=['get'])
    def pages(self, request, pk=None):
        """PDF 페이지를 이미지(base64)로 반환 — 프론트 오버레이용"""
        document = self.get_object()
        pdf_bytes = None

        if document.pdf_file:
            document.pdf_file.open('rb')
            pdf_bytes = document.pdf_file.read()
            document.pdf_file.close()
        elif document.file and document.file.name.lower().endswith('.pdf'):
            document.file.open('rb')
            pdf_bytes = document.file.read()
            document.file.close()
        elif document.file and document.file.name.lower().endswith('.docx'):
            pdf_bytes = convert_docx_to_pdf(document.file)
            if pdf_bytes:
                filename = f'{document.document_type.lower()}_{document.id}.pdf'
                document.pdf_file.save(filename, ContentFile(pdf_bytes), save=True)

        if not pdf_bytes:
            return Response({'error': 'No PDF available'}, status=404)

        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            pages_data = []
            for page in doc:
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")
                img_b64 = base64.b64encode(img_bytes).decode('utf-8')
                pages_data.append({
                    'image': f'data:image/png;base64,{img_b64}',
                    'width': pix.width,
                    'height': pix.height,
                })
            num_pages = len(pages_data)
            doc.close()

            # 서명 영역 좌표 반환 (저장된 zones 또는 기본값)
            zones = document.sign_zones if document.sign_zones else get_default_zones(num_pages)

            return Response({
                'pages': pages_data,
                'total': num_pages,
                'sign_zones': zones,
            })
        except Exception as e:
            return Response({'error': f'Failed to render pages: {str(e)}'}, status=500)

    @action(detail=True, methods=['get'])
    def download(self, request, pk=None):
        """서명 완료된 문서 다운로드 (PDF)"""
        document = self.get_object()

        if not document.is_signed:
            return Response(
                {'error': 'Document must be signed first'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # PDF 우선, 없으면 원본
        if document.pdf_file:
            content = document.pdf_file.read()
        elif document.file:
            content = document.file.read()
        else:
            return Response({'error': 'No file found'}, status=404)

        document.download_count += 1
        document.save(update_fields=['download_count'])

        response = HttpResponse(content, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{document.title}.pdf"'
        return response

    @action(detail=False, methods=['get'], url_path='my-documents')
    def my_documents(self, request):
        """현재 로그인 유저의 문서 목록"""
        docs = EmployeeDocument.objects.filter(
            employee=request.user.profile
        ).select_related('signed_by__user', 'uploaded_by__user').order_by('-uploaded_at')
        serializer = self.get_serializer(docs, many=True)
        return Response(serializer.data)


class RosterViewSet(viewsets.ModelViewSet):
    queryset = Roster.objects.all()
    serializer_class = RosterSerializer
    permission_classes = [IsAuthenticated, IsManager]
    filter_backends = [OrganizationFilterBackend]

    def create(self, request, *args, **kwargs):
        """Upsert: 같은 user+date가 있으면 업데이트"""
        user_id = request.data.get('user')
        date_val = request.data.get('date')
        existing = Roster.objects.filter(user_id=user_id, date=date_val).first()
        if existing:
            serializer = self.get_serializer(existing, data=request.data, partial=True)
            serializer.is_valid(raise_exception=True)
            serializer.save(organization=self._resolve_org())
            return Response(serializer.data, status=status.HTTP_200_OK)
        return super().create(request, *args, **kwargs)

    def _resolve_org(self):
        profile = self.request.user.profile
        if profile.role in ['CEO', 'HQ', 'REGIONAL_MANAGER', 'SENIOR_MANAGER']:
            store_id = self.request.query_params.get('store_id')
            if store_id:
                from users.models import Organization
                try:
                    return Organization.objects.get(id=store_id)
                except Organization.DoesNotExist:
                    pass
        return profile.organization

    def perform_create(self, serializer):
        serializer.save(organization=self._resolve_org())

    def update(self, request, *args, **kwargs):
        """Update: user/date 변경 시 충돌하면 기존 것 삭제 후 이동"""
        instance = self.get_object()
        new_user = request.data.get('user', instance.user_id)
        new_date = request.data.get('date', instance.date)
        # 다른 셀로 이동할 때 기존 로스터 충돌 처리
        if (int(new_user) != instance.user_id or str(new_date) != str(instance.date)):
            conflict = Roster.objects.filter(user_id=new_user, date=new_date).exclude(id=instance.id).first()
            if conflict:
                conflict.delete()
        return super().update(request, *args, **kwargs)

    def get_queryset(self):
        """사용자의 조직에 해당하는 스케줄만 조회"""
        queryset = super().get_queryset()
        return queryset.select_related('user__user', 'organization').order_by('-date')

    @action(detail=False, methods=['post'], url_path='bulk-break')
    def bulk_break(self, request):
        """Update break_minutes for all unpublished rosters in a week."""
        break_minutes = request.data.get('break_minutes', 30)
        date_str = request.data.get('date')
        target_date = timezone.localdate() if not date_str else timezone.datetime.strptime(date_str, '%Y-%m-%d').date()
        week_start = target_date - timedelta(days=target_date.weekday())
        week_end = week_start + timedelta(days=6)
        org = self._resolve_org()

        updated = Roster.objects.filter(
            organization=org,
            date__range=[week_start, week_end],
            is_confirmed=False,
        ).update(break_minutes=break_minutes)

        return Response({'updated': updated, 'break_minutes': break_minutes})

    @action(detail=False, methods=['get'])
    def weekly(self, request):
        """주간 스케줄 조회"""
        try:
            date_str = request.query_params.get('date')
            target_date = timezone.localdate() if not date_str else timezone.datetime.strptime(date_str, '%Y-%m-%d').date()

            # 주의 시작일 (월요일)
            week_start = target_date - timedelta(days=target_date.weekday())
            week_end = week_start + timedelta(days=6)

            rosters = self.filter_queryset(self.get_queryset()).filter(
                date__range=[week_start, week_end]
            )

            serializer = self.get_serializer(rosters, many=True)
            return Response({
                'week_start': week_start,
                'week_end': week_end,
                'rosters': serializer.data
            }, status=status.HTTP_200_OK)

        except Exception as e:
            return Response(
                {'error': str(e)},
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=False, methods=['get'])
    def monthly(self, request):
        """월간 스케줄 조회"""
        try:
            date_str = request.query_params.get('date')
            if date_str:
                target_date = timezone.datetime.strptime(date_str, '%Y-%m-%d').date()
            else:
                target_date = timezone.localdate()

            month_start = target_date.replace(day=1)
            if month_start.month == 12:
                month_end = month_start.replace(year=month_start.year + 1, month=1, day=1) - timedelta(days=1)
            else:
                month_end = month_start.replace(month=month_start.month + 1) - timedelta(days=1)

            rosters = self.filter_queryset(self.get_queryset()).filter(
                date__range=[month_start, month_end]
            )

            serializer = self.get_serializer(rosters, many=True)
            return Response({
                'month_start': month_start,
                'month_end': month_end,
                'rosters': serializer.data
            }, status=status.HTTP_200_OK)

        except Exception as e:
            return Response(
                {'error': str(e)},
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=False, methods=['get'], url_path='my-roster',
            permission_classes=[IsAuthenticated])
    def my_roster(self, request):
        """Employee's own roster: today, tomorrow, and weekly"""
        profile = request.user.profile
        today = timezone.localdate()  # NZ date, not UTC
        tomorrow = today + timedelta(days=1)

        today_roster = Roster.objects.filter(user=profile, date=today).first()
        tomorrow_roster = Roster.objects.filter(user=profile, date=tomorrow).first()

        week_start = today - timedelta(days=today.weekday())
        week_end = week_start + timedelta(days=6)
        weekly = Roster.objects.filter(
            user=profile, date__range=[week_start, week_end]
        ).order_by('date')

        return Response({
            'today': RosterSerializer(today_roster).data if today_roster else None,
            'tomorrow': RosterSerializer(tomorrow_roster).data if tomorrow_roster else None,
            'week_start': week_start,
            'week_end': week_end,
            'weekly': RosterSerializer(weekly, many=True).data,
        })

    @action(detail=False, methods=['post'], url_path='copy-week')
    def copy_week(self, request):
        """Copy last week's roster to a target week"""
        try:
            source_date = request.data.get('source_date')  # any date in source week
            target_date = request.data.get('target_date')  # any date in target week

            if not source_date or not target_date:
                return Response({'error': 'source_date and target_date required'}, status=status.HTTP_400_BAD_REQUEST)

            from datetime import datetime
            src = datetime.strptime(source_date, '%Y-%m-%d').date()
            tgt = datetime.strptime(target_date, '%Y-%m-%d').date()

            src_monday = src - timedelta(days=src.weekday())
            tgt_monday = tgt - timedelta(days=tgt.weekday())
            day_offset = (tgt_monday - src_monday).days

            org = request.user.profile.organization
            source_rosters = Roster.objects.filter(
                organization=org,
                date__range=[src_monday, src_monday + timedelta(days=6)]
            )

            created = 0
            skipped = 0
            for r in source_rosters:
                new_date = r.date + timedelta(days=day_offset)
                _, was_created = Roster.objects.get_or_create(
                    user=r.user,
                    date=new_date,
                    defaults={
                        'organization': org,
                        'shift_start': r.shift_start,
                        'shift_end': r.shift_end,
                        'shift_name': r.shift_name,
                        'shift_color': r.shift_color,
                        'shift_template': r.shift_template,
                        'break_minutes': r.break_minutes,
                        'is_confirmed': False,
                    }
                )
                if was_created:
                    created += 1
                else:
                    skipped += 1

            return Response({
                'created': created,
                'skipped': skipped,
                'message': f'{created} shifts copied, {skipped} skipped (already exist)'
            })

        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=False, methods=['post'], url_path='bulk-delete')
    def bulk_delete(self, request):
        """Bulk delete unconfirmed rosters by IDs"""
        ids = request.data.get('ids', [])
        if not ids:
            return Response({'detail': 'ids list is required.'}, status=status.HTTP_400_BAD_REQUEST)

        org = request.user.profile.organization
        deleted_count, _ = Roster.objects.filter(
            id__in=ids,
            organization=org,
            is_confirmed=False,
        ).delete()

        return Response({'deleted': deleted_count}, status=status.HTTP_200_OK)


class ShiftTemplateViewSet(viewsets.ModelViewSet):
    """시프트 템플릿 CRUD"""
    queryset = ShiftTemplate.objects.all()
    serializer_class = ShiftTemplateSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = None

    def get_queryset(self):
        org = get_target_org(self.request)
        return ShiftTemplate.objects.filter(organization=org, is_active=True)

    def create(self, request, *args, **kwargs):
        from django.db import IntegrityError
        try:
            return super().create(request, *args, **kwargs)
        except IntegrityError:
            return Response(
                {'name': ['A template with this name already exists.']},
                status=status.HTTP_400_BAD_REQUEST
            )

    def perform_create(self, serializer):
        serializer.save(organization=get_target_org(self.request))

    def perform_update(self, serializer):
        instance = serializer.save()
        # Sync color/name/times to all unpublished rosters using this template
        Roster.objects.filter(
            shift_template=instance,
            is_confirmed=False,
        ).update(
            shift_name=instance.name,
            shift_color=instance.color,
            shift_start=instance.start_time,
            shift_end=instance.end_time,
            break_minutes=instance.break_minutes,
        )


class TimesheetViewSet(viewsets.ModelViewSet):
    queryset = Timesheet.objects.all()
    serializer_class = TimesheetSerializer
    permission_classes = [IsAuthenticated]
    filter_backends = [OrganizationFilterBackend]

    def get_permissions(self):
        employee_actions = ['today', 'clock_in', 'clock_out', 'start_break', 'end_break']
        if self.action in employee_actions:
            return [IsAuthenticated()]
        return [IsAuthenticated(), IsManager()]

    def get_queryset(self):
        queryset = super().get_queryset()
        return queryset.select_related('user__user', 'organization').order_by('-date')

    @action(detail=True, methods=['post'])
    def approve(self, request, pk=None):
        timesheet = self.get_object()
        if timesheet.is_approved:
            return Response({'error': 'Already approved.'}, status=status.HTTP_400_BAD_REQUEST)
        timesheet.is_approved = True
        timesheet.save()
        return Response({'message': 'Timesheet approved.'})

    @action(detail=True, methods=['post'], url_path='approve-overtime')
    def approve_overtime(self, request, pk=None):
        """Manager approves overtime for a timesheet"""
        timesheet = self.get_object()
        if not timesheet.is_overtime:
            return Response({'error': 'No overtime to approve.'}, status=status.HTTP_400_BAD_REQUEST)
        if timesheet.overtime_approved:
            return Response({'error': 'Overtime already approved.'}, status=status.HTTP_400_BAD_REQUEST)
        timesheet.overtime_approved = True
        timesheet.overtime_approved_by = request.user.profile
        timesheet.save()
        return Response({'message': 'Overtime approved.'})

    @action(detail=False, methods=['get'])
    def weekly(self, request):
        try:
            date_str = request.query_params.get('date')
            target_date = timezone.localdate() if not date_str else timezone.datetime.strptime(date_str, '%Y-%m-%d').date()
            week_start = target_date - timedelta(days=target_date.weekday())
            week_end = week_start + timedelta(days=6)
            timesheets = self.filter_queryset(self.get_queryset()).filter(
                date__range=[week_start, week_end]
            )
            serializer = self.get_serializer(timesheets, many=True)
            return Response({
                'week_start': week_start,
                'week_end': week_end,
                'timesheets': serializer.data
            })
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    # ---- Employee Time Clock Actions ----

    @action(detail=False, methods=['get'])
    def today(self, request):
        """Return current user's today timesheet + status"""
        profile = request.user.profile
        today = timezone.localdate()
        try:
            ts = Timesheet.objects.get(user=profile, date=today)
        except Timesheet.DoesNotExist:
            return Response({'status': 'NOT_CLOCKED_IN', 'timesheet': None})

        if ts.check_out:
            clock_status = 'CLOCKED_OUT'
        elif ts.break_start and not ts.break_end:
            clock_status = 'ON_BREAK'
        else:
            clock_status = 'WORKING'

        serializer = self.get_serializer(ts)

        # Calculate scheduled hours for overtime detection
        roster = Roster.objects.filter(user=profile, date=today).first()
        scheduled_hours = roster.hours if roster else None

        return Response({
            'status': clock_status,
            'timesheet': serializer.data,
            'scheduled_hours': scheduled_hours,
        })

    @action(detail=False, methods=['post'], url_path='clock_in')
    def clock_in(self, request):
        profile = request.user.profile
        today = timezone.localdate()
        if Timesheet.objects.filter(user=profile, date=today).exists():
            return Response({'error': 'Already clocked in today.'}, status=status.HTTP_400_BAD_REQUEST)
        ts = Timesheet.objects.create(
            organization=profile.organization,
            user=profile,
            date=today,
            check_in=timezone.now(),
        )
        serializer = self.get_serializer(ts)
        return Response({'status': 'WORKING', 'timesheet': serializer.data}, status=status.HTTP_201_CREATED)

    @action(detail=False, methods=['post'], url_path='clock_out')
    def clock_out(self, request):
        profile = request.user.profile
        today = timezone.localdate()
        try:
            ts = Timesheet.objects.get(user=profile, date=today)
        except Timesheet.DoesNotExist:
            return Response({'error': 'No active timesheet.'}, status=status.HTTP_400_BAD_REQUEST)

        if ts.check_out:
            return Response({'error': 'Already clocked out.'}, status=status.HTTP_400_BAD_REQUEST)

        # End active break if any
        if ts.break_start and not ts.break_end:
            ts.break_end = timezone.now()
            delta = (ts.break_end - ts.break_start).total_seconds() / 60
            ts.total_break_minutes += int(delta)

        # Validation: incomplete tasks
        incomplete_tasks = Task.objects.filter(
            assigned_to=profile, due_date=today
        ).exclude(status='COMPLETED')

        # Validation: FCP daily checklist
        # Only check if: 1) employee has can_safety_tasks permission AND 2) store has active templates
        from safety.models import DailyChecklistResponse, SafetyChecklistTemplate
        fcp_done = True  # default: no blocker
        if profile.can_safety_tasks:
            has_fcp_templates = SafetyChecklistTemplate.objects.filter(
                organization=profile.organization,
                is_active=True,
            ).exists()
            if has_fcp_templates:
                fcp_done = DailyChecklistResponse.objects.filter(
                    organization=profile.organization,
                    date=today,
                    completed_by=profile,
                    is_completed=True,
                ).exists()

        blockers = []
        if incomplete_tasks.exists():
            blockers.append({
                'type': 'tasks',
                'message': f'{incomplete_tasks.count()} task(s) not completed.',
                'items': list(incomplete_tasks.values('id', 'title', 'status')),
            })
        if not fcp_done:
            blockers.append({
                'type': 'fcp',
                'message': 'FCP daily checklist not completed.',
            })

        if blockers:
            return Response({
                'status': 'BLOCKED',
                'blockers': blockers,
                'timesheet': self.get_serializer(ts).data,
            }, status=status.HTTP_409_CONFLICT)

        # Overtime check
        overtime_reason = request.data.get('overtime_reason', '')
        roster = Roster.objects.filter(user=profile, date=today).first()
        now = timezone.now()

        # Calculate actual worked hours at checkout
        total_seconds = (now - ts.check_in).total_seconds()
        break_seconds = ts.total_break_minutes * 60
        net_hours = (total_seconds - break_seconds) / 3600

        is_overtime = False
        if roster and net_hours > roster.hours:
            is_overtime = True

        if is_overtime and not overtime_reason:
            return Response({
                'status': 'OVERTIME_REQUIRED',
                'message': 'Overtime detected. Please provide a reason.',
                'scheduled_hours': roster.hours if roster else None,
                'actual_hours': round(net_hours, 2),
                'timesheet': self.get_serializer(ts).data,
            }, status=status.HTTP_409_CONFLICT)

        ts.check_out = now
        if is_overtime:
            ts.is_overtime = True
            ts.overtime_reason = overtime_reason
        ts.save()

        return Response({
            'status': 'CLOCKED_OUT',
            'timesheet': self.get_serializer(ts).data,
        })

    @action(detail=False, methods=['post'], url_path='break_start')
    def start_break(self, request):
        profile = request.user.profile
        today = timezone.localdate()
        try:
            ts = Timesheet.objects.get(user=profile, date=today)
        except Timesheet.DoesNotExist:
            return Response({'error': 'Not clocked in.'}, status=status.HTTP_400_BAD_REQUEST)

        if ts.check_out:
            return Response({'error': 'Already clocked out.'}, status=status.HTTP_400_BAD_REQUEST)
        if ts.break_start and not ts.break_end:
            return Response({'error': 'Already on break.'}, status=status.HTTP_400_BAD_REQUEST)

        ts.break_start = timezone.now()
        ts.break_end = None
        ts.save()
        return Response({'status': 'ON_BREAK', 'timesheet': self.get_serializer(ts).data})

    @action(detail=False, methods=['post'], url_path='break_end')
    def end_break(self, request):
        profile = request.user.profile
        today = timezone.localdate()
        try:
            ts = Timesheet.objects.get(user=profile, date=today)
        except Timesheet.DoesNotExist:
            return Response({'error': 'Not clocked in.'}, status=status.HTTP_400_BAD_REQUEST)

        if not ts.break_start or ts.break_end:
            return Response({'error': 'Not on break.'}, status=status.HTTP_400_BAD_REQUEST)

        ts.break_end = timezone.now()
        delta = (ts.break_end - ts.break_start).total_seconds() / 60
        ts.total_break_minutes += int(delta)
        ts.save()
        return Response({'status': 'WORKING', 'timesheet': self.get_serializer(ts).data})


class TaskViewSet(viewsets.ModelViewSet):
    """
    업무 할당 관리 ViewSet
    """
    queryset = Task.objects.all()
    serializer_class = TaskSerializer
    permission_classes = [IsAuthenticated]
    filter_backends = [OrganizationFilterBackend]

    def get_queryset(self):
        queryset = super().get_queryset()
        queryset = queryset.select_related(
            'assigned_to__user', 'assigned_by__user', 'organization'
        )
        # Filter by status if provided
        status_filter = self.request.query_params.get('status')
        if status_filter:
            queryset = queryset.filter(status=status_filter)
        return queryset

    def perform_create(self, serializer):
        profile = self.request.user.profile
        serializer.save(
            assigned_by=profile,
            organization=profile.organization
        )

    @action(detail=True, methods=['post'])
    def complete(self, request, pk=None):
        task = self.get_object()
        task.status = 'COMPLETED'
        task.completed_at = timezone.now()
        task.save()
        return Response({'message': 'Task completed.'}, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'])
    def start(self, request, pk=None):
        task = self.get_object()
        task.status = 'IN_PROGRESS'
        task.save()
        return Response({'message': 'Task started.'}, status=status.HTTP_200_OK)


class TeamViewSet(viewsets.ViewSet):
    """팀 멤버 관리 (같은 조직 직원들)"""
    permission_classes = [IsAuthenticated]

    def _get_hourly_rate(self, profile):
        salary = Salary.objects.filter(user=profile, is_active=True).first()
        return salary.hourly_rate if salary else None

    def _get_org(self, request):
        """CEO/HQ는 store_id로 조직 전환 가능"""
        profile = request.user.profile
        if profile.role in ['CEO', 'HQ']:
            store_id = request.query_params.get('store_id')
            if store_id:
                from users.models import Organization
                try:
                    return Organization.objects.get(id=store_id)
                except Organization.DoesNotExist:
                    pass
        return profile.organization

    def list(self, request):
        profile = request.user.profile
        org = self._get_org(request)
        status_filter = request.query_params.get('status', 'ACTIVE')
        search = request.query_params.get('search', '').strip()

        team = UserProfile.objects.filter(
            organization=org,
        ).select_related('user')

        if status_filter and status_filter != 'ALL':
            team = team.filter(employment_status=status_filter)

        if search:
            from django.db.models import Q
            team = team.filter(
                Q(user__first_name__icontains=search) |
                Q(user__last_name__icontains=search) |
                Q(user__username__icontains=search) |
                Q(user__email__icontains=search) |
                Q(employee_id__icontains=search) |
                Q(phone__icontains=search)
            )

        data = []
        for p in team:
            data.append({
                'id': p.id,
                'profile_id': p.id,
                'user_id': p.user.id,
                'name': p.user.get_full_name() or p.user.username,
                'employee_id': p.employee_id,
                'role': p.role,
                'role_display': dict(ROLE_CHOICES).get(p.role, p.role),
                'job_title': p.job_title,
                'job_title_display': dict(JOB_TITLE_CHOICES).get(p.job_title) if p.job_title else None,
                'work_type': p.work_type,
                'work_type_display': dict(WORK_TYPE_CHOICES).get(p.work_type, p.work_type),
                'employment_status': p.employment_status,
                'date_of_joining': p.date_of_joining,
                'phone': p.phone,
                'hourly_rate': self._get_hourly_rate(p),
            })
        return Response(data, status=status.HTTP_200_OK)

    def retrieve(self, request, pk=None):
        """직원 상세 정보"""
        org = self._get_org(request)
        try:
            member = UserProfile.objects.select_related('user').get(
                id=pk, organization=org
            )
        except UserProfile.DoesNotExist:
            return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)

        # Salary history
        salaries = Salary.objects.filter(user=member).order_by('-effective_from')
        salary_history = [
            {
                'id': s.id,
                'hourly_rate': str(s.hourly_rate),
                'overtime_multiplier': str(s.overtime_multiplier),
                'effective_from': s.effective_from,
                'effective_to': s.effective_to,
                'is_active': s.is_active,
            }
            for s in salaries
        ]

        # Documents
        documents = EmployeeDocument.objects.filter(employee=member).order_by('-uploaded_at')
        doc_serializer = EmployeeDocumentSerializer(documents, many=True, context={'request': request})

        # Onboarding status
        onboarding_status = None
        try:
            ob = member.onboarding
            onboarding_status = ob.status
        except Onboarding.DoesNotExist:
            pass

        data = {
            'id': member.id,
            'user_id': member.user.id,
            'name': member.user.get_full_name() or member.user.username,
            'first_name': member.user.first_name,
            'last_name': member.user.last_name,
            'email': member.user.email,
            'employee_id': member.employee_id,
            'role': member.role,
            'role_display': dict(ROLE_CHOICES).get(member.role, member.role),
            'job_title': member.job_title,
            'job_title_display': dict(JOB_TITLE_CHOICES).get(member.job_title) if member.job_title else None,
            'work_type': member.work_type,
            'work_type_display': dict(WORK_TYPE_CHOICES).get(member.work_type, member.work_type),
            'employment_status': member.employment_status,
            'date_of_joining': member.date_of_joining,
            'date_of_birth': member.date_of_birth,
            'phone': member.phone,
            'kiwisaver_status': member.kiwisaver_status,
            'kiwisaver_rate': member.kiwisaver_rate,
            'hourly_rate': self._get_hourly_rate(member),
            'salary_history': salary_history,
            'documents': doc_serializer.data,
            'onboarding_status': onboarding_status,
            'can_daily_close': member.can_daily_close,
            'can_safety_tasks': member.can_safety_tasks,
            'housing_support': member.housing_support,
            'housing_amount': str(member.housing_amount),
            'transport_support': member.transport_support,
            'transport_amount': str(member.transport_amount),
            'annual_salary': str(member.annual_salary),
        }
        return Response(data)

    @action(detail=True, methods=['post'], url_path='update-salary')
    def update_salary(self, request, pk=None):
        """시급 변경"""
        profile = request.user.profile
        try:
            member = UserProfile.objects.get(id=pk, organization=profile.organization)
        except UserProfile.DoesNotExist:
            return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)

        hourly_rate = request.data.get('hourly_rate')
        if not hourly_rate:
            return Response({'error': 'hourly_rate is required'}, status=status.HTTP_400_BAD_REQUEST)

        # Support custom effective_from date (for future salary changes)
        effective_from_str = request.data.get('effective_from')
        if effective_from_str:
            from datetime import date as date_cls
            try:
                effective_date = date_cls.fromisoformat(effective_from_str)
            except (ValueError, TypeError):
                return Response({'error': 'Invalid date format. Use YYYY-MM-DD'}, status=status.HTTP_400_BAD_REQUEST)
        else:
            effective_date = timezone.localdate()

        with transaction.atomic():
            # Deactivate current salary (set end date to day before new rate)
            Salary.objects.filter(user=member, is_active=True).update(
                is_active=False,
                effective_to=effective_date - timedelta(days=1),
            )
            # Create new salary
            new_salary = Salary.objects.create(
                organization=profile.organization,
                user=member,
                hourly_rate=hourly_rate,
                effective_from=effective_date,
                is_active=True,
            )

        return Response({
            'message': 'Salary updated',
            'hourly_rate': str(new_salary.hourly_rate),
            'effective_from': new_salary.effective_from,
        })

    @action(detail=True, methods=['post'], url_path='update-permissions')
    def update_permissions(self, request, pk=None):
        """직원 태스크 권한 + 프로필 업데이트 (매니저 전용)"""
        profile = request.user.profile
        MANAGER_ROLES = ['MANAGER', 'SENIOR_MANAGER', 'REGIONAL_MANAGER', 'HQ', 'CEO']
        if profile.role not in MANAGER_ROLES:
            return Response({'error': 'Permission denied'}, status=status.HTTP_403_FORBIDDEN)

        org = self._get_org(request)
        try:
            member = UserProfile.objects.get(id=pk, organization=org)
        except UserProfile.DoesNotExist:
            return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)

        updated = []
        for field in ['can_daily_close', 'can_safety_tasks', 'housing_support', 'transport_support']:
            if field in request.data:
                setattr(member, field, bool(request.data[field]))
                updated.append(field)

        # Decimal fields
        from decimal import Decimal, InvalidOperation
        for field in ['housing_amount', 'transport_amount', 'annual_salary']:
            if field in request.data:
                try:
                    setattr(member, field, Decimal(str(request.data[field])))
                    updated.append(field)
                except (InvalidOperation, ValueError):
                    pass

        if 'job_title' in request.data:
            member.job_title = request.data['job_title'] or None
            updated.append('job_title')

        if 'work_type' in request.data:
            valid_types = [c[0] for c in WORK_TYPE_CHOICES]
            if request.data['work_type'] in valid_types:
                member.work_type = request.data['work_type']
                updated.append('work_type')

        if 'kiwisaver_status' in request.data:
            valid_statuses = ['NOT_ENROLLED', 'ENROLLED', 'OPTED_OUT']
            if request.data['kiwisaver_status'] in valid_statuses:
                member.kiwisaver_status = request.data['kiwisaver_status']
                updated.append('kiwisaver_status')

        if 'kiwisaver_rate' in request.data:
            valid_rates = ['3%', '4%', '6%', '8%', '10%']
            if request.data['kiwisaver_rate'] in valid_rates:
                member.kiwisaver_rate = request.data['kiwisaver_rate']
                updated.append('kiwisaver_rate')

        if updated:
            member.save(update_fields=updated + ['updated_at'])

        return Response({
            'message': 'Updated',
            'can_daily_close': member.can_daily_close,
            'can_safety_tasks': member.can_safety_tasks,
            'housing_support': member.housing_support,
            'housing_amount': str(member.housing_amount),
            'transport_support': member.transport_support,
            'transport_amount': str(member.transport_amount),
            'annual_salary': str(member.annual_salary),
            'work_type': member.work_type,
            'work_type_display': dict(WORK_TYPE_CHOICES).get(member.work_type, member.work_type),
            'job_title': member.job_title,
            'job_title_display': dict(JOB_TITLE_CHOICES).get(member.job_title) if member.job_title else None,
            'kiwisaver_status': member.kiwisaver_status,
            'kiwisaver_rate': member.kiwisaver_rate,
        })

    @action(detail=True, methods=['post'], url_path='reset-password')
    def reset_password(self, request, pk=None):
        """매니저가 직원 비밀번호 리셋 (임시 비밀번호 생성 후 이메일 발송)"""
        profile = request.user.profile
        MANAGER_ROLES = ['MANAGER', 'SENIOR_MANAGER', 'REGIONAL_MANAGER', 'HQ', 'CEO']
        if profile.role not in MANAGER_ROLES:
            return Response({'error': 'Permission denied'}, status=status.HTTP_403_FORBIDDEN)

        org = self._get_org(request)
        try:
            member = UserProfile.objects.select_related('user').get(id=pk, organization=org)
        except UserProfile.DoesNotExist:
            return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)

        user = member.user
        # Generate temporary password
        temp_password = User.objects.make_random_password(length=10)
        user.set_password(temp_password)
        user.is_active = True  # Activate account if not yet active
        user.save()

        # Send email with new password
        try:
            login_url = 'https://oneops.co.nz/login'
            html_message = render_to_string('hr/password_reset_by_manager.html', {
                'first_name': user.first_name or user.username,
                'email': user.email,
                'temp_password': temp_password,
                'login_url': login_url,
                'manager_name': request.user.get_full_name() or request.user.username,
            })
            send_mail(
                subject='[Oneops] Your Password Has Been Reset',
                message=f'Hi {user.first_name}, your password has been reset. '
                        f'Login at {login_url} with email: {user.email} and temporary password: {temp_password}. '
                        f'Please change your password after logging in.',
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[user.email],
                html_message=html_message,
                fail_silently=False,
            )
            logger.info('Password reset email sent to %s by manager %s', user.email, request.user.email)
        except Exception as e:
            logger.error('Failed to send password reset email: %s', e)
            # Password is already changed, so return success but warn about email
            return Response({
                'message': f'Password reset successfully. Temporary password: {temp_password} (Email delivery failed, please share manually)',
                'temp_password': temp_password,
                'email_sent': False,
            })

        return Response({
            'message': f'Password reset email sent to {user.email}',
            'email_sent': True,
        })

    @action(detail=True, methods=['get'])
    def documents(self, request, pk=None):
        """직원 문서 목록"""
        profile = request.user.profile
        try:
            member = UserProfile.objects.get(id=pk, organization=profile.organization)
        except UserProfile.DoesNotExist:
            return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)

        docs = EmployeeDocument.objects.filter(employee=member).order_by('-uploaded_at')
        serializer = EmployeeDocumentSerializer(docs, many=True, context={'request': request})
        return Response(serializer.data)

    @action(detail=True, methods=['post'], url_path='upload-document')
    def upload_document(self, request, pk=None):
        """직원에게 문서 업로드"""
        profile = request.user.profile
        try:
            member = UserProfile.objects.get(id=pk, organization=profile.organization)
        except UserProfile.DoesNotExist:
            return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)

        serializer = EmployeeDocumentSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        serializer.save(
            employee=member,
            organization=profile.organization,
            uploaded_by=profile,
        )
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=['get'], url_path='employee-file')
    def employee_file(self, request, pk=None):
        """직원 파일 - 모든 기록 통합 조회 (NZ compliance)"""
        org = self._get_org(request)
        try:
            member = UserProfile.objects.select_related('user').get(id=pk, organization=org)
        except UserProfile.DoesNotExist:
            return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)

        # Employee basic info
        employee_info = {
            'id': member.id,
            'name': member.user.get_full_name() or member.user.username,
            'email': member.user.email,
            'employee_id': member.employee_id,
            'role': member.role,
            'job_title': member.job_title,
            'job_title_display': dict(JOB_TITLE_CHOICES).get(member.job_title) if member.job_title else None,
            'work_type': member.work_type,
            'work_type_display': dict(WORK_TYPE_CHOICES).get(member.work_type, member.work_type),
            'employment_status': member.employment_status,
            'date_of_joining': member.date_of_joining,
            'phone': member.phone,
        }

        # 1. Employment Agreements
        employment_types = ['CONTRACT', 'JOB_OFFER', 'JOB_DESCRIPTION', 'VARIATION']
        employment_docs = EmployeeDocument.objects.filter(
            employee=member, document_type__in=employment_types
        ).order_by('-uploaded_at')

        # 2. Tax Declarations
        ir330s = IR330Declaration.objects.filter(employee=member).order_by('-created_at')

        # 3. Wage & Time Records
        from payroll.models import Salary, PaySlip
        salaries = Salary.objects.filter(user=member).order_by('-effective_from')
        recent_payslips = PaySlip.objects.filter(user=member).order_by('-pay_period_end')[:6]
        timesheet_count = Timesheet.objects.filter(user=member).count()
        timesheet_range = Timesheet.objects.filter(user=member).order_by('date').values_list('date', flat=True)
        timesheet_dates = list(timesheet_range[:1]) + list(timesheet_range.reverse()[:1])

        # 4. Leave Records
        from payroll.models import LeaveRequest, LeaveBalance
        leave_balances = LeaveBalance.objects.filter(user=member)
        leave_requests = LeaveRequest.objects.filter(user=member).order_by('-created_at')[:20]

        # 5. Training & Certifications
        cert_docs = EmployeeDocument.objects.filter(
            employee=member, document_type__in=['CERTIFICATE']
        ).order_by('-uploaded_at')
        try:
            onboarding = member.onboarding
            training_tasks = OnboardingTask.objects.filter(
                onboarding=onboarding, step_type='TRAINING'
            ).order_by('order')
        except Onboarding.DoesNotExist:
            training_tasks = OnboardingTask.objects.none()

        # 6. Disciplinary Records
        disciplinary = DisciplinaryRecord.objects.filter(employee=member).order_by('-date')

        # 7. Performance Reviews
        performance = PerformanceReview.objects.filter(employee=member).order_by('-review_period_end')

        # 8. Health & Safety (Workplace Accidents)
        accidents = WorkplaceAccident.objects.filter(employee=member).order_by('-date')

        # 9. Other Documents (visa, ID, medical, police vetting, etc.)
        other_doc_types = ['VISA', 'ID_DOCUMENT', 'MEDICAL', 'POLICE_VET', 'OTHER']
        other_docs = EmployeeDocument.objects.filter(
            employee=member, document_type__in=other_doc_types
        ).order_by('-uploaded_at')

        # 10. File Notes
        notes = EmployeeNote.objects.filter(employee=member).order_by('-date')

        # 11. Inquiries
        inquiries = Inquiry.objects.filter(employee=member).order_by('-created_at')

        # 12. Resignation
        resignations = ResignationRequest.objects.filter(employee=member).order_by('-created_at')

        # 13. Onboarding
        onboarding_data = None
        try:
            ob = member.onboarding
            onboarding_data = {
                'status': ob.status,
                'completed_percentage': ob.completed_percentage,
                'completed_at': ob.completed_at,
                'created_at': ob.created_at,
            }
        except Onboarding.DoesNotExist:
            pass

        # Build timeline (chronological events from all categories)
        timeline = []

        for doc in employment_docs:
            timeline.append({
                'date': str(doc.uploaded_at.date()) if doc.uploaded_at else None,
                'type': 'document',
                'category': 'employment',
                'title': f'{doc.get_document_type_display()}: {doc.title}',
                'detail': f'Signed: {"Yes" if doc.is_signed else "No"}',
                'id': doc.id,
            })

        for ir in ir330s:
            timeline.append({
                'date': str(ir.created_at.date()),
                'type': 'tax',
                'category': 'tax',
                'title': f'IR330 Tax Declaration - {ir.get_tax_code_display()}',
                'detail': f'IRD: {ir.ird_number}',
                'id': ir.id,
            })

        for d in disciplinary:
            timeline.append({
                'date': str(d.date),
                'type': 'disciplinary',
                'category': 'disciplinary',
                'title': f'{d.get_record_type_display()}: {d.subject}',
                'detail': d.description[:100],
                'id': d.id,
            })

        for p in performance:
            timeline.append({
                'date': str(p.review_period_end),
                'type': 'performance',
                'category': 'performance',
                'title': f'Performance Review - {p.get_overall_rating_display()}',
                'detail': f'{p.review_period_start} ~ {p.review_period_end}',
                'id': p.id,
            })

        for a in accidents:
            timeline.append({
                'date': str(a.date),
                'type': 'accident',
                'category': 'health_safety',
                'title': f'Workplace Accident - {a.get_injury_type_display()}',
                'detail': a.description[:100],
                'id': a.id,
            })

        for n in notes:
            timeline.append({
                'date': str(n.date),
                'type': 'note',
                'category': 'notes',
                'title': f'{n.get_category_display()}: {n.subject}',
                'detail': n.content[:100],
                'id': n.id,
            })

        for inq in inquiries:
            timeline.append({
                'date': str(inq.created_at.date()),
                'type': 'inquiry',
                'category': 'inquiries',
                'title': f'Inquiry: {inq.subject}',
                'detail': f'Status: {inq.get_status_display()}',
                'id': inq.id,
            })

        for r in resignations:
            timeline.append({
                'date': str(r.created_at.date()),
                'type': 'resignation',
                'category': 'resignation',
                'title': f'Resignation - {r.get_status_display()}',
                'detail': f'Last day: {r.requested_last_day}',
                'id': r.id,
            })

        # Sort timeline by date descending
        timeline.sort(key=lambda x: x.get('date') or '', reverse=True)

        # Serialize categories
        response_data = {
            'employee': employee_info,
            'categories': {
                'employment': {
                    'label': 'Employment Agreements',
                    'count': employment_docs.count(),
                    'items': EmployeeDocumentSerializer(employment_docs, many=True, context={'request': request}).data,
                },
                'tax': {
                    'label': 'Tax Declarations (IR330)',
                    'count': ir330s.count(),
                    'items': IR330DeclarationSerializer(ir330s, many=True).data,
                },
                'wages_time': {
                    'label': 'Wage & Time Records',
                    'count': timesheet_count,
                    'salary_history': [{
                        'id': s.id,
                        'hourly_rate': str(s.hourly_rate),
                        'effective_from': s.effective_from,
                        'effective_to': s.effective_to,
                        'is_active': s.is_active,
                    } for s in salaries],
                    'recent_payslips': [{
                        'id': p.id,
                        'pay_period_start': p.pay_period_start,
                        'pay_period_end': p.pay_period_end,
                        'gross_pay': str(p.gross_pay),
                        'net_pay': str(p.net_pay),
                        'status': p.status,
                    } for p in recent_payslips],
                    'timesheet_summary': {
                        'total_records': timesheet_count,
                        'earliest': str(timesheet_dates[0]) if timesheet_dates else None,
                        'latest': str(timesheet_dates[-1]) if len(timesheet_dates) > 1 else None,
                    },
                },
                'leave': {
                    'label': 'Leave Records',
                    'balances': [{
                        'leave_type': lb.leave_type,
                        'total_days': float(lb.total_days),
                        'used_days': float(lb.used_days),
                        'remaining_days': float(lb.remaining_days),
                    } for lb in leave_balances],
                    'recent_requests': [{
                        'id': lr.id,
                        'leave_type': lr.leave_type,
                        'start_date': lr.start_date,
                        'end_date': lr.end_date,
                        'status': lr.status,
                        'reason': lr.reason,
                    } for lr in leave_requests],
                },
                'training': {
                    'label': 'Training & Certifications',
                    'certificates': EmployeeDocumentSerializer(cert_docs, many=True, context={'request': request}).data,
                    'training_tasks': [{
                        'id': t.id,
                        'title': t.title,
                        'is_completed': t.is_completed,
                        'completed_at': t.completed_at,
                    } for t in training_tasks],
                },
                'disciplinary': {
                    'label': 'Disciplinary Records',
                    'count': disciplinary.count(),
                    'items': DisciplinaryRecordSerializer(disciplinary, many=True).data,
                },
                'performance': {
                    'label': 'Performance Reviews',
                    'count': performance.count(),
                    'items': PerformanceReviewSerializer(performance, many=True).data,
                },
                'health_safety': {
                    'label': 'Health & Safety',
                    'count': accidents.count(),
                    'items': WorkplaceAccidentSerializer(accidents, many=True).data,
                },
                'documents': {
                    'label': 'Other Documents',
                    'count': other_docs.count(),
                    'items': EmployeeDocumentSerializer(other_docs, many=True, context={'request': request}).data,
                },
                'notes': {
                    'label': 'File Notes',
                    'count': notes.count(),
                    'items': EmployeeNoteSerializer(notes, many=True).data,
                },
                'inquiries': {
                    'label': 'Inquiries',
                    'count': inquiries.count(),
                    'items': InquirySerializer(inquiries, many=True).data,
                },
                'resignation': {
                    'label': 'Resignation',
                    'items': ResignationRequestSerializer(resignations, many=True).data,
                },
                'onboarding': {
                    'label': 'Onboarding',
                    'data': onboarding_data,
                },
            },
            'timeline': timeline,
        }

        return Response(response_data)


def send_invite_email(invite, temp_password, org_name):
    """초대 이메일 발송"""
    try:
        login_url = 'https://oneops.co.nz/login'
        html_message = render_to_string('hr/invite_email.html', {
            'first_name': invite.first_name,
            'organization': org_name,
            'email': invite.email,
            'temp_password': temp_password,
            'login_url': login_url,
        })
        send_mail(
            subject=f'[Oneops] Welcome to {org_name} - Your Account is Ready',
            message=f'Hi {invite.first_name}, you have been invited to join {org_name} on Oneops. '
                    f'Login at {login_url} with email: {invite.email} and password: {temp_password}',
            from_email=settings.DEFAULT_FROM_EMAIL,
            recipient_list=[invite.email],
            html_message=html_message,
            fail_silently=False,
        )
        logger.info(f'Invitation email sent to {invite.email}')
        return True
    except Exception as e:
        logger.error(f'Failed to send invitation email to {invite.email}: {e}')
        return False


class EmployeeInviteViewSet(viewsets.ModelViewSet):
    """직원 초대 관리"""
    serializer_class = EmployeeInviteSerializer
    permission_classes = [IsAuthenticated, IsManager]
    pagination_class = None

    def get_queryset(self):
        org = get_target_org(self.request)
        return EmployeeInvite.objects.filter(organization=org).select_related(
            'invited_by__user', 'accepted_by__user'
        )

    def create(self, request, *args, **kwargs):
        """초대 생성 → User + UserProfile + Salary + Onboarding 자동 생성 (is_active=False)"""
        from hr.utils import generate_temp_password, generate_username_from_email

        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)

        profile = request.user.profile
        org = get_target_org(request)
        invite_data = serializer.validated_data

        email = invite_data['email']

        # 중복 이메일 체크
        if User.objects.filter(email=email).exists():
            return Response({'error': 'A user with this email already exists'},
                            status=status.HTTP_400_BAD_REQUEST)

        temp_password = generate_temp_password()
        username = generate_username_from_email(email)

        with transaction.atomic():
            # 1. Save invite
            invite = serializer.save(
                organization=org,
                invited_by=profile,
                expires_at=timezone.now() + timedelta(days=7),
            )

            # 2. Create User (is_active=False until email verification via invite link)
            user = User.objects.create_user(
                username=username,
                password=temp_password,
                email=email,
                first_name=invite.first_name,
                last_name=invite.last_name,
                is_active=False,
            )

            # 3. Create UserProfile
            employee_id = f"EMP-{org.id}-{user.id}"
            user_profile = UserProfile.objects.create(
                user=user,
                employee_id=employee_id,
                role=invite.role,
                organization=org,
                job_title=invite.job_title,
                work_type=invite.work_type,
                date_of_joining=timezone.localdate(),
                phone='',
            )

            # 4. Create Salary
            Salary.objects.create(
                organization=org,
                user=user_profile,
                hourly_rate=invite.hourly_rate,
                effective_from=timezone.localdate(),
                is_active=True,
            )

            # 5. Create Onboarding + Tasks
            onboarding = Onboarding.objects.create(
                organization=org,
                employee=user_profile,
                status='IN_PROGRESS',
                assigned_to=profile,
            )

            steps = [
                {'order': 1, 'step_type': 'PERSONAL_INFO', 'title': 'Personal Information'},
                {'order': 2, 'step_type': 'BANK_ACCOUNT', 'title': 'Bank Account'},
                {'order': 3, 'step_type': 'IR330', 'title': 'IR330 Tax Declaration'},
            ]

            # Document templates → auto-generate documents
            docx_context = {
                'employee_name': f'{invite.first_name} {invite.last_name}'.strip(),
                'employee_first_name': invite.first_name,
                'employee_last_name': invite.last_name or '',
                'employee_email': invite.email,
                'hourly_rate': str(invite.hourly_rate),
                'holiday_rate': str(round(float(invite.hourly_rate) * 0.08, 2)),
                'gross_rate': str(round(float(invite.hourly_rate) * 1.08, 2)),
                'hours': str(invite.min_hours or 30),
                'job_title': invite.job_title or '',
                'position_title': invite.job_title or '',
                'work_type': invite.get_work_type_display(),
                'start_date': timezone.now().strftime('%d/%m/%Y'),
                'commencement_date': invite.commencement_date.strftime('%d/%m/%Y') if invite.commencement_date else 'TBC',
                'work_location': invite.work_location or org.address or '',
                'min_hours': str(invite.min_hours or 30),
                'max_hours': str(invite.max_hours or 50),
                'reporting_to': invite.reporting_to or 'Director/Management',
                'company_name': org.name,
                'company_address': org.address or '',
                'company_phone': org.phone or '',
                'company_email': org.email or '',
                'company_ird': org.ird_number or '',
            }

            order_num = 4
            for doc_type in ['CONTRACT', 'JOB_OFFER', 'JOB_DESCRIPTION']:
                filters = {
                    'organization': org,
                    'document_type': doc_type,
                    'is_active': True,
                }
                if doc_type == 'CONTRACT':
                    filters['work_type'] = invite.work_type
                elif doc_type == 'JOB_DESCRIPTION':
                    filters['job_title'] = invite.job_title
                template = DocumentTemplate.objects.filter(**filters).first()
                if template:
                    doc = EmployeeDocument.objects.create(
                        onboarding=onboarding,
                        employee=user_profile,
                        organization=org,
                        document_type=doc_type,
                        title=template.title,
                        uploaded_by=profile,
                    )

                    if template.file and template.file.name.endswith('.docx'):
                        try:
                            filled = fill_docx_template(template.file, docx_context)
                            filename = f'{doc_type.lower()}_{invite.first_name}_{doc.id}.docx'
                            doc.file.save(filename, ContentFile(filled.read()), save=True)

                            try:
                                pdf_bytes = convert_docx_to_pdf(doc.file)
                                if pdf_bytes:
                                    cleaned_pdf, zones = extract_and_clean_sign_zones(pdf_bytes)
                                    pdf_name = f'{doc_type.lower()}_{doc.id}.pdf'
                                    doc.pdf_file.save(pdf_name, ContentFile(cleaned_pdf), save=False)
                                    doc.sign_zones = zones
                                    doc.save()
                            except Exception:
                                pass
                        except Exception:
                            doc.file = template.file
                            doc.save()
                    else:
                        doc.file = template.file
                        doc.save()

                    steps.append({
                        'order': order_num,
                        'step_type': 'DOCUMENT_SIGN',
                        'title': f'Sign {template.get_document_type_display()}',
                        'related_document_id': doc.id,
                    })
                    order_num += 1

            # File uploads
            steps.append({'order': order_num, 'step_type': 'FILE_UPLOAD', 'title': 'Upload Visa', 'upload_label': 'Visa'})
            order_num += 1
            steps.append({'order': order_num, 'step_type': 'FILE_UPLOAD', 'title': 'Upload Resume', 'upload_label': 'Resume'})
            order_num += 1

            # Training modules
            for module_type in ['SAFETY', 'FCP', 'HAZARD']:
                module = TrainingModule.objects.filter(
                    organization=org,
                    module_type=module_type,
                    is_active=True,
                ).first()
                if module:
                    steps.append({
                        'order': order_num,
                        'step_type': 'TRAINING',
                        'title': module.title,
                        'related_training_id': module.id,
                    })
                    order_num += 1

            for step in steps:
                OnboardingTask.objects.create(onboarding=onboarding, **step)

            # Link invite to the created profile
            invite.accepted_by = user_profile
            invite.save()

        # Send invitation email
        email_sent = send_invite_email(invite, temp_password, org.name)

        # Return invite data + generated credentials
        response_data = self.get_serializer(invite).data
        response_data['generated_credentials'] = {
            'email': email,
            'password': temp_password,
        }
        response_data['email_sent'] = email_sent
        return Response(response_data, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=['post'])
    def resend(self, request, pk=None):
        """초대 재발급 (새 code + 새 만료) + 이메일 재발송"""
        from hr.utils import generate_temp_password
        invite = self.get_object()
        invite.invite_code = uuid.uuid4()
        invite.status = 'PENDING'
        invite.expires_at = timezone.now() + timedelta(days=7)
        invite.save()

        # Reset password and resend email
        temp_password = generate_temp_password()
        user = User.objects.filter(email=invite.email).first()
        if user:
            user.set_password(temp_password)
            user.save()

        org = request.user.profile.organization
        email_sent = send_invite_email(invite, temp_password, org.name)

        response_data = self.get_serializer(invite).data
        response_data['email_sent'] = email_sent
        return Response(response_data)


class AcceptInviteView(APIView):
    """초대 수락 → 이메일 인증 (계정 활성화) (Public - 로그인 불필요)"""
    permission_classes = [AllowAny]

    def get(self, request):
        """초대 코드 검증 → 초대 정보 반환"""
        code = request.query_params.get('code')
        if not code:
            return Response({'error': 'code is required'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            invite = EmployeeInvite.objects.select_related('organization').get(invite_code=code)
        except EmployeeInvite.DoesNotExist:
            return Response({'error': 'Invalid invite code'}, status=status.HTTP_404_NOT_FOUND)

        if invite.status == 'ACCEPTED':
            return Response({'error': 'Invite already accepted'}, status=status.HTTP_400_BAD_REQUEST)
        if invite.is_expired:
            return Response({'error': 'Invite expired'}, status=status.HTTP_400_BAD_REQUEST)

        return Response({
            'first_name': invite.first_name,
            'last_name': invite.last_name,
            'email': invite.email,
            'organization_name': invite.organization.name,
            'job_title': invite.job_title or '',
            'role': invite.get_role_display(),
        })

    def post(self, request):
        """초대 수락 → 계정 활성화 (이메일 인증 완료) → JWT 발급"""
        code = request.data.get('invite_code')

        if not code:
            return Response({'error': 'invite_code is required'},
                            status=status.HTTP_400_BAD_REQUEST)

        try:
            invite = EmployeeInvite.objects.select_related(
                'organization', 'accepted_by__user'
            ).get(invite_code=code)
        except EmployeeInvite.DoesNotExist:
            return Response({'error': 'Invalid invite code'}, status=status.HTTP_404_NOT_FOUND)

        if invite.status == 'ACCEPTED':
            return Response({'error': 'Invite already accepted'}, status=status.HTTP_400_BAD_REQUEST)
        if invite.is_expired:
            return Response({'error': 'Invite expired'}, status=status.HTTP_400_BAD_REQUEST)

        # Find the user created during invite
        user_profile = invite.accepted_by
        if not user_profile:
            return Response({'error': 'No account linked to this invite'},
                            status=status.HTTP_400_BAD_REQUEST)

        user = user_profile.user

        with transaction.atomic():
            # Activate the user account (email verification complete)
            user.is_active = True
            user.save(update_fields=['is_active'])

            # Update invite status
            invite.status = 'ACCEPTED'
            invite.save()

        # Generate JWT tokens for auto-login
        from rest_framework_simplejwt.tokens import RefreshToken
        refresh = RefreshToken.for_user(user)

        # Find onboarding record
        onboarding = Onboarding.objects.filter(
            employee=user_profile, status='IN_PROGRESS'
        ).first()

        return Response({
            'message': 'Account activated successfully',
            'user_id': user.id,
            'profile_id': user_profile.id,
            'access': str(refresh.access_token),
            'refresh': str(refresh),
            'onboarding_id': onboarding.id if onboarding else None,
        }, status=status.HTTP_200_OK)


class DocumentTemplateViewSet(viewsets.ModelViewSet):
    """문서 템플릿 관리 (Store Settings용)"""
    serializer_class = DocumentTemplateSerializer
    permission_classes = [IsAuthenticated, IsManager]
    pagination_class = None

    def get_queryset(self):
        org = self.request.user.profile.organization
        return DocumentTemplate.objects.filter(organization=org)

    def perform_create(self, serializer):
        serializer.save(organization=get_target_org(self.request))

    # Known placeholders that are auto-filled during invite creation
    KNOWN_PLACEHOLDERS = {
        'employee_name': 'Employee full name',
        'employee_first_name': 'Employee first name',
        'employee_last_name': 'Employee last name',
        'employee_email': 'Employee email',
        'hourly_rate': 'Hourly rate',
        'holiday_rate': 'Holiday rate (8% of hourly)',
        'gross_rate': 'Gross rate (hourly + holiday)',
        'hours': 'Weekly hours',
        'job_title': 'Job title',
        'position_title': 'Position title',
        'work_type': 'Work type (Full Time, Part Time, etc.)',
        'start_date': 'Start date',
        'commencement_date': 'Commencement date',
        'work_location': 'Work location',
        'min_hours': 'Minimum hours',
        'max_hours': 'Maximum hours',
        'reporting_to': 'Reports to',
        'company_name': 'Company name',
        'company_address': 'Company address',
        'company_phone': 'Company phone',
        'company_email': 'Company email',
        'company_ird': 'Company IRD number',
    }

    @action(detail=False, methods=['post'])
    def extract_placeholders(self, request):
        """Extract {{placeholder}} tokens from an uploaded DOCX file."""
        file = request.FILES.get('file')
        if not file:
            return Response({'error': 'No file provided'}, status=status.HTTP_400_BAD_REQUEST)

        if not file.name.endswith('.docx'):
            return Response({'error': 'Only .docx files are supported'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            from docx import Document as DocxDocument
            import re

            doc = DocxDocument(file)
            placeholder_re = re.compile(r'\{\{(\w+)\}\}')
            found = set()

            # Extract from body paragraphs
            for p in doc.paragraphs:
                full_text = ''.join(run.text for run in p.runs)
                found.update(placeholder_re.findall(full_text))

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            full_text = ''.join(run.text for run in p.runs)
                            found.update(placeholder_re.findall(full_text))

            # Extract from headers/footers
            for section in doc.sections:
                for part in [section.header, section.footer,
                             section.first_page_header, section.first_page_footer]:
                    if part and part.is_linked_to_previous is False:
                        for p in part.paragraphs:
                            full_text = ''.join(run.text for run in p.runs)
                            found.update(placeholder_re.findall(full_text))

            known = []
            unknown = []
            for ph in sorted(found):
                if ph in self.KNOWN_PLACEHOLDERS:
                    known.append({'key': ph, 'description': self.KNOWN_PLACEHOLDERS[ph]})
                else:
                    unknown.append({'key': ph, 'description': None})

            return Response({
                'filename': file.name,
                'total': len(found),
                'known': known,
                'unknown': unknown,
                'known_keys': list(self.KNOWN_PLACEHOLDERS.keys()),
            })
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=False, methods=['post'])
    def fix_placeholders(self, request):
        """Replace incorrect placeholders in a DOCX and return the fixed file."""
        file = request.FILES.get('file')
        if not file:
            return Response({'error': 'No file provided'}, status=status.HTTP_400_BAD_REQUEST)

        import json
        mappings_raw = request.data.get('mappings', '{}')
        if isinstance(mappings_raw, str):
            try:
                mappings = json.loads(mappings_raw)
            except json.JSONDecodeError:
                return Response({'error': 'Invalid mappings JSON'}, status=status.HTTP_400_BAD_REQUEST)
        else:
            mappings = mappings_raw

        if not mappings:
            return Response({'error': 'No mappings provided'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            from docx import Document as DocxDocument
            import re
            from io import BytesIO

            doc = DocxDocument(file)

            def replace_in_paragraph(paragraph):
                full_text = ''.join(run.text for run in paragraph.runs)
                new_text = full_text
                for old_key, new_key in mappings.items():
                    new_text = new_text.replace(f'{{{{{old_key}}}}}', f'{{{{{new_key}}}}}')
                if new_text != full_text and paragraph.runs:
                    paragraph.runs[0].text = new_text
                    for run in paragraph.runs[1:]:
                        run.text = ''

            # Body paragraphs
            for p in doc.paragraphs:
                replace_in_paragraph(p)

            # Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_in_paragraph(p)

            # Headers/footers
            for section in doc.sections:
                for part in [section.header, section.footer,
                             section.first_page_header, section.first_page_footer]:
                    if part and part.is_linked_to_previous is False:
                        for p in part.paragraphs:
                            replace_in_paragraph(p)

            output = BytesIO()
            doc.save(output)
            output.seek(0)

            from django.http import FileResponse
            response = FileResponse(
                output,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="fixed_{file.name}"'
            return response
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)


class TrainingModuleViewSet(viewsets.ModelViewSet):
    """교육 모듈 관리 (Store Settings용)"""
    serializer_class = TrainingModuleSerializer
    permission_classes = [IsAuthenticated, IsManager]
    pagination_class = None

    def get_queryset(self):
        org = self.request.user.profile.organization
        return TrainingModule.objects.filter(organization=org)

    def perform_create(self, serializer):
        serializer.save(organization=get_target_org(self.request))


class IR330ViewSet(viewsets.ModelViewSet):
    """IR330 Tax Declaration 관리"""
    serializer_class = IR330DeclarationSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = None

    def get_queryset(self):
        return IR330Declaration.objects.filter(
            employee__organization=self.request.user.profile.organization
        )

    def perform_create(self, serializer):
        # Handle signature base64 image
        signature_data = self.request.data.get('signature')
        save_kwargs = {'employee': self.request.user.profile, 'signed_at': timezone.now()}

        if signature_data and ';base64,' in str(signature_data):
            try:
                fmt, imgstr = signature_data.split(';base64,')
                ext = fmt.split('/')[-1] if '/' in fmt else 'png'
                sig_file = ContentFile(
                    base64.b64decode(imgstr),
                    name=f'ir330_sig_{self.request.user.id}.{ext}',
                )
                save_kwargs['signature'] = sig_file
            except Exception:
                pass

        serializer.save(**save_kwargs)


class SaveBankAccountView(generics.GenericAPIView):
    """온보딩 시 뱅크 어카운트 저장"""
    permission_classes = [IsAuthenticated]

    def post(self, request):
        bank_account = request.data.get('bank_account', '').strip()
        if not bank_account:
            return Response({'error': 'Bank account is required'}, status=400)
        profile = request.user.profile
        profile.bank_account = bank_account
        profile.save(update_fields=['bank_account'])
        return Response({'status': 'ok', 'bank_account': bank_account})


class InquiryViewSet(viewsets.ModelViewSet):
    """직원 문의 ViewSet — 직원은 본인 문의만, 매니저는 조직 전체"""
    serializer_class = InquirySerializer
    permission_classes = [IsAuthenticated]

    MANAGER_ROLES = ['MANAGER', 'SENIOR_MANAGER', 'REGIONAL_MANAGER', 'HQ', 'CEO']

    def get_queryset(self):
        user = self.request.user.profile
        if user.role in self.MANAGER_ROLES:
            return Inquiry.objects.filter(
                organization=user.organization
            ).select_related('employee__user', 'replied_by__user')
        return Inquiry.objects.filter(
            employee=user
        ).select_related('employee__user', 'replied_by__user')

    def perform_create(self, serializer):
        serializer.save(
            employee=self.request.user.profile,
            organization=self.request.user.profile.organization,
        )

    @action(detail=True, methods=['post'])
    def reply(self, request, pk=None):
        """매니저가 답변"""
        inquiry = self.get_object()
        reply_text = request.data.get('reply', '').strip()
        if not reply_text:
            return Response({'error': 'Reply is required'}, status=400)

        inquiry.reply = reply_text
        inquiry.replied_by = request.user.profile
        inquiry.replied_at = timezone.now()
        inquiry.status = 'REPLIED'
        inquiry.save()
        return Response(InquirySerializer(inquiry).data)

    @action(detail=True, methods=['post'])
    def close(self, request, pk=None):
        """문의 종료"""
        inquiry = self.get_object()
        inquiry.status = 'CLOSED'
        inquiry.save(update_fields=['status'])
        return Response(InquirySerializer(inquiry).data)


class ResignationRequestViewSet(viewsets.ModelViewSet):
    """퇴직 신청 ViewSet — 직원은 본인만, 매니저는 조직 전체"""
    serializer_class = ResignationRequestSerializer
    permission_classes = [IsAuthenticated]

    MANAGER_ROLES = ['MANAGER', 'SENIOR_MANAGER', 'REGIONAL_MANAGER', 'HQ', 'CEO']

    def get_queryset(self):
        user = self.request.user.profile
        if user.role in self.MANAGER_ROLES:
            return ResignationRequest.objects.filter(
                organization=user.organization
            ).select_related('employee__user', 'confirmed_by__user')
        return ResignationRequest.objects.filter(
            employee=user
        ).select_related('employee__user', 'confirmed_by__user')

    def perform_create(self, serializer):
        profile = self.request.user.profile
        # Calculate notice period based on work type
        work_type = profile.work_type
        if work_type in ('FULL_TIME', 'SALARY', 'VISA_FULL_TIME'):
            notice_weeks = 2
        else:  # PART_TIME, CASUAL
            notice_weeks = 1

        from datetime import date, timedelta
        today = date.today()
        earliest = today + timedelta(weeks=notice_weeks)

        serializer.save(
            employee=profile,
            organization=profile.organization,
            notice_period_weeks=notice_weeks,
            earliest_last_day=earliest,
        )

    @action(detail=True, methods=['post'])
    def confirm(self, request, pk=None):
        """매니저가 퇴직 확인 (last day 네고 가능)"""
        resignation = self.get_object()
        if resignation.status != 'PENDING':
            return Response({'error': 'Only pending requests can be confirmed'}, status=400)

        confirmed_last_day = request.data.get('confirmed_last_day')
        manager_notes = request.data.get('manager_notes', '')

        if confirmed_last_day:
            resignation.confirmed_last_day = confirmed_last_day
        else:
            resignation.confirmed_last_day = resignation.requested_last_day

        resignation.manager_notes = manager_notes
        resignation.confirmed_by = request.user.profile
        resignation.confirmed_at = timezone.now()
        resignation.status = 'CONFIRMED'
        resignation.save()
        return Response(ResignationRequestSerializer(resignation).data)

    @action(detail=True, methods=['post'])
    def withdraw(self, request, pk=None):
        """직원이 퇴직 철회"""
        resignation = self.get_object()
        if resignation.status != 'PENDING':
            return Response({'error': 'Only pending requests can be withdrawn'}, status=400)
        if resignation.employee != request.user.profile:
            return Response({'error': 'Permission denied'}, status=403)
        resignation.status = 'WITHDRAWN'
        resignation.save(update_fields=['status'])
        return Response(ResignationRequestSerializer(resignation).data)


class DisciplinaryRecordViewSet(viewsets.ModelViewSet):
    """징계 기록 CRUD — 매니저 전용"""
    serializer_class = DisciplinaryRecordSerializer
    permission_classes = [IsAuthenticated]

    MANAGER_ROLES = ['MANAGER', 'SENIOR_MANAGER', 'REGIONAL_MANAGER', 'HQ', 'CEO']

    def get_queryset(self):
        user = self.request.user.profile
        qs = DisciplinaryRecord.objects.filter(
            organization=user.organization
        ).select_related('employee__user', 'issued_by__user')
        employee_id = self.request.query_params.get('employee')
        if employee_id:
            qs = qs.filter(employee_id=employee_id)
        return qs

    def perform_create(self, serializer):
        profile = self.request.user.profile
        serializer.save(
            organization=profile.organization,
            issued_by=profile,
        )


class PerformanceReviewViewSet(viewsets.ModelViewSet):
    """성과 평가 CRUD — 매니저 전용"""
    serializer_class = PerformanceReviewSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user.profile
        qs = PerformanceReview.objects.filter(
            organization=user.organization
        ).select_related('employee__user', 'reviewer__user')
        employee_id = self.request.query_params.get('employee')
        if employee_id:
            qs = qs.filter(employee_id=employee_id)
        return qs

    def perform_create(self, serializer):
        profile = self.request.user.profile
        serializer.save(
            organization=profile.organization,
            reviewer=profile,
        )


class WorkplaceAccidentViewSet(viewsets.ModelViewSet):
    """산업재해/사고 기록 CRUD — 매니저 전용"""
    serializer_class = WorkplaceAccidentSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user.profile
        qs = WorkplaceAccident.objects.filter(
            organization=user.organization
        ).select_related('employee__user', 'reported_by__user')
        employee_id = self.request.query_params.get('employee')
        if employee_id:
            qs = qs.filter(employee_id=employee_id)
        return qs

    def perform_create(self, serializer):
        profile = self.request.user.profile
        serializer.save(
            organization=profile.organization,
            reported_by=profile,
        )


class EmployeeNoteViewSet(viewsets.ModelViewSet):
    """직원 파일 노트 CRUD — 매니저 전용"""
    serializer_class = EmployeeNoteSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user.profile
        qs = EmployeeNote.objects.filter(
            organization=user.organization
        ).select_related('employee__user', 'created_by__user')
        employee_id = self.request.query_params.get('employee')
        if employee_id:
            qs = qs.filter(employee_id=employee_id)
        return qs

    def perform_create(self, serializer):
        profile = self.request.user.profile
        serializer.save(
            organization=profile.organization,
            created_by=profile,
        )
