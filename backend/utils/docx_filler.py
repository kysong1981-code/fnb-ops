"""
DOCX Template Filler Utility
Replaces {{placeholder}} tokens in DOCX files with actual data.
Handles Word's run-splitting issue by working at paragraph level.
"""
import re
from io import BytesIO
from copy import deepcopy
from docx import Document


PLACEHOLDER_RE = re.compile(r'\{\{(\w+)\}\}')


def _replace_in_paragraph(paragraph, context):
    """
    Replace placeholders in a paragraph.
    Word often splits {{placeholder}} across multiple runs, e.g.:
      Run1: "{{employee"  Run2: "_name}}"
    Solution: join all runs, do replacement on full text,
    then put result in first run and clear the rest.
    """
    full_text = ''.join(run.text for run in paragraph.runs)
    if not PLACEHOLDER_RE.search(full_text):
        return

    # Replace all placeholders
    new_text = PLACEHOLDER_RE.sub(
        lambda m: str(context.get(m.group(1), m.group(0))),
        full_text,
    )

    if new_text == full_text:
        return

    # Put replaced text into first run, clear the rest
    # Preserve the formatting of the first run
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''


def _replace_in_table(table, context):
    """Replace placeholders in all cells of a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, context)
            # Handle nested tables
            for nested_table in cell.tables:
                _replace_in_table(nested_table, context)


def fill_docx_template(template_file, context):
    """
    Fill a DOCX template with context data.

    Args:
        template_file: Django FileField or file-like object
        context: dict of {placeholder_name: value}
                 e.g. {'employee_name': 'John Doe', 'hourly_rate': '25.00'}

    Returns:
        BytesIO object containing the filled DOCX (ready for Django file save)
    """
    # Read the template
    template_file.seek(0)
    doc = Document(template_file)

    # Replace in body paragraphs
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, context)

    # Replace in tables
    for table in doc.tables:
        _replace_in_table(table, context)

    # Replace in headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.is_linked_to_previous is False:
                for paragraph in header.paragraphs:
                    _replace_in_paragraph(paragraph, context)
                for table in header.tables:
                    _replace_in_table(table, context)

        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                for paragraph in footer.paragraphs:
                    _replace_in_paragraph(paragraph, context)
                for table in footer.tables:
                    _replace_in_table(table, context)

    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
