"""Convert DOCX to PDF using python-docx + fpdf2 (pure Python, no external tools)."""
import io

from docx import Document
from docx.oxml.ns import qn
from fpdf import FPDF


def _clean_text(text):
    """Clean text for Helvetica (Latin-1) compatibility."""
    replacements = {
        '\u2013': '-',   # en dash
        '\u2014': '--',  # em dash
        '\u2018': "'",   # left single quote
        '\u2019': "'",   # right single quote
        '\u201c': '"',   # left double quote
        '\u201d': '"',   # right double quote
        '\u2022': '*',   # bullet
        '\u2026': '...',  # ellipsis
        '\u00a0': ' ',   # non-breaking space
        '\u2010': '-',   # hyphen
        '\u2011': '-',   # non-breaking hyphen
        '\u2012': '-',   # figure dash
        '\u00b7': '*',   # middle dot
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.encode('latin-1', errors='replace').decode('latin-1')


def convert_docx_to_pdf(docx_content):
    """
    Convert DOCX file content to PDF bytes.

    Args:
        docx_content: bytes, file-like object, or Django FieldFile

    Returns:
        bytes: PDF content, or None if conversion fails
    """
    try:
        # Handle different input types
        if hasattr(docx_content, 'open'):
            # Django FieldFile
            with docx_content.open('rb') as f:
                stream = io.BytesIO(f.read())
        elif hasattr(docx_content, 'read'):
            pos = docx_content.tell() if hasattr(docx_content, 'tell') else 0
            stream = io.BytesIO(docx_content.read())
            if hasattr(docx_content, 'seek'):
                docx_content.seek(pos)
        elif isinstance(docx_content, bytes):
            stream = io.BytesIO(docx_content)
        else:
            return None

        doc = Document(stream)

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()
        pdf.set_font('Helvetica', size=10)

        pw = pdf.w - pdf.l_margin - pdf.r_margin  # usable page width

        # Iterate body elements in document order
        paragraphs = list(doc.paragraphs)
        tables = list(doc.tables)
        p_idx = 0
        t_idx = 0

        for element in doc.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

            if tag == 'p' and p_idx < len(paragraphs):
                _render_paragraph(pdf, paragraphs[p_idx], pw)
                p_idx += 1
            elif tag == 'tbl' and t_idx < len(tables):
                _render_table(pdf, tables[t_idx], pw)
                t_idx += 1

        return bytes(pdf.output())
    except Exception:
        return None


def _render_paragraph(pdf, para, pw):
    """Render a single paragraph to PDF."""
    text = _clean_text(para.text.strip())
    style = para.style.name if para.style else ''

    if not text:
        pdf.ln(3)
        return

    # Title / Heading 1
    if 'Title' in style or 'Heading 1' in style:
        pdf.ln(5)
        pdf.set_font('Helvetica', 'B', 16)
        pdf.multi_cell(pw, 8, text)
        pdf.ln(3)
        pdf.set_font('Helvetica', size=10)
        return

    # Heading 2
    if 'Heading 2' in style:
        pdf.ln(4)
        pdf.set_font('Helvetica', 'B', 13)
        pdf.multi_cell(pw, 7, text)
        pdf.ln(2)
        pdf.set_font('Helvetica', size=10)
        return

    # Heading 3+
    if 'Heading' in style:
        pdf.ln(3)
        pdf.set_font('Helvetica', 'B', 11)
        pdf.multi_cell(pw, 6, text)
        pdf.ln(2)
        pdf.set_font('Helvetica', size=10)
        return

    # Normal paragraph — check if all runs are bold
    runs = [r for r in para.runs if r.text.strip()]
    is_bold = runs and all(r.bold for r in runs)

    if is_bold:
        pdf.set_font('Helvetica', 'B', 10)

    pdf.multi_cell(pw, 5, text)
    pdf.ln(1)

    if is_bold:
        pdf.set_font('Helvetica', size=10)


def _render_table(pdf, table, pw):
    """Render a table to PDF."""
    pdf.ln(3)

    if not table.rows:
        return

    num_cols = max(len(row.cells) for row in table.rows)
    if num_cols == 0:
        return

    col_w = pw / num_cols

    for i, row in enumerate(table.rows):
        # Deduplicate merged cells
        seen = set()
        cells = []
        for cell in row.cells:
            cell_id = id(cell)
            if cell_id not in seen:
                seen.add(cell_id)
                cells.append(_clean_text(cell.text.strip()))

        # Pad to num_cols
        while len(cells) < num_cols:
            cells.append('')

        # Calculate row height
        max_lines = 1
        for text in cells:
            if text:
                chars_per_line = max(1, int(col_w * 0.45))
                lines = max(1, (len(text) // chars_per_line) + 1)
                max_lines = max(max_lines, lines)

        row_h = max(6, 5 * max_lines)

        # Page break check
        if pdf.get_y() + row_h > pdf.h - pdf.b_margin:
            pdf.add_page()

        y0 = pdf.get_y()
        x0 = pdf.get_x()

        for j, text in enumerate(cells[:num_cols]):
            x = x0 + j * col_w
            # Draw cell border
            pdf.rect(x, y0, col_w, row_h)

            # Header row bold
            if i == 0:
                pdf.set_font('Helvetica', 'B', 9)
            else:
                pdf.set_font('Helvetica', size=9)

            pdf.set_xy(x + 1, y0 + 1)
            pdf.multi_cell(col_w - 2, 4, text[:200])  # truncate very long text

        pdf.set_xy(x0, y0 + row_h)

    pdf.ln(3)
    pdf.set_font('Helvetica', size=10)
